import os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class GenerateReport:
    def __init__(self, dict):
        self.data_dict = dict
        self.document = Document()


    def getVulnDescribe(self, table, name, num):
        cells1 = table.rows[0].cells
        cells1[0].paragraphs[0].add_run('漏洞名称: ' + '\n' + name)
        cells1[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="3D9140"/>'.format(nsdecls('w'))))
        cells2 = table.rows[1].cells
        cells2[0].paragraphs[0].add_run('数量: ' + str(num))
        cells2[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w'))))
        cells3 = table.rows[2].cells
        cells3[0].paragraphs[0].add_run('详细信息: ')
        cells3[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="7FFFD4"/>'.format(nsdecls('w'))))
        cells5 = table.rows[4].cells
        cells5[0].paragraphs[0].add_run('修复建议: ')
        cells5[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="7FFFD4"/>'.format(nsdecls('w'))))


    def getTable(self, table, data_list):
        for index, vuln in enumerate(data_list):
            cells1 = table.rows[index * 3 + 6].cells
            cells1[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="00FFFF"/>'.format(nsdecls('w'))))
            cells1[0].paragraphs[0].add_run('风险点' + str(index + 1))
            cells2 = table.rows[index * 3 + 7].cells
            cells2[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FAEBD7"/>'.format(nsdecls('w'))))
            cells2[0].paragraphs[0].add_run('触发行：' + '\n' + str(vuln[0]) + '行')
            cells3 = table.rows[index * 3 + 8].cells
            if len(vuln) > 2:
                for i in range(1, len(vuln)):
                    if (i == 1):
                        cells3[0].paragraphs[0].add_run('调用路径1: ')
                    else:
                        cells3[0].paragraphs[0].add_run('\n' + '调用路径' + str(i) + ': ')
                    for num, trace in enumerate(vuln[i]):
                        cells3[0].paragraphs[0].add_run('\n' + str(trace))
            else:
                cells3[0].paragraphs[0].add_run('调用路径：')
                for num, trace in enumerate(vuln[1]):
                    cells3[0].paragraphs[0].add_run('\n' + str(trace))


    def start(self):
        for index, key in enumerate(self.data_dict.keys()):
            self.document.add_heading(str(index+1) + '. ' + str(key), level=1)
            table = self.document.add_table(rows=len(self.data_dict[key]) * 3 + 6, cols=1, style="Table Grid")
            self.getVulnDescribe(table, key, len(self.data_dict[key]))
            self.getTable(table, self.data_dict[key])

        if (os.path.exists('./Report')):
            self.document.save('./Report/test.docx')
        else:
            os.mkdir('./Report')
            self.document.save('./Report/test.docx')